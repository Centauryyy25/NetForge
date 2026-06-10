#!/usr/bin/env python3
"""Generate BAB III .docx for SI YBY NET project."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
DIAG = ROOT / "diagrams"
OUT = ROOT / "BAB_III_Perancangan_Sistem.docx"

# ---------- helpers ----------
doc = Document()

# Page setup: A4, margin kiri 4cm, lainnya 3cm
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Cm(4)
section.right_margin = Cm(3)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.first_line_indent = Cm(0)

for h_name, h_size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12), ("Heading 4", 12)]:
    s = doc.styles[h_name]
    s.font.name = "Times New Roman"
    s.font.size = Pt(h_size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    s.paragraph_format.line_spacing = 1.5


def p(text="", bold=False, italic=False, align="justify", indent_first=True, size=12):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.5
    if indent_first:
        par.paragraph_format.first_line_indent = Cm(1.27)
    par.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if text:
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
    return par


def h(text, level=1):
    par = doc.add_heading(text, level=level)
    par.paragraph_format.line_spacing = 1.5
    for run in par.runs:
        run.font.name = "Times New Roman"
    return par


def caption(text, align="center"):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing = 1.0
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(text)
    run.italic = False
    run.bold = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_image(path: Path, width_in=6.0, cap=None):
    if not path.exists():
        p(f"[Gambar tidak tersedia: {path.name}]", italic=True, align="center", indent_first=False)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    if cap:
        caption(cap)


def code_block(text):
    par = doc.add_paragraph()
    par.paragraph_format.line_spacing = 1.15
    par.paragraph_format.left_indent = Cm(0.5)
    par.paragraph_format.space_after = Pt(6)
    run = par.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    # shaded background
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)


def bullet(text, level=0):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.line_spacing = 1.15
    par.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    run = par.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def make_table(headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(htxt)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            para.paragraph_format.line_spacing = 1.15
            run = para.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    # add small space after
    doc.add_paragraph()
    return t


def page_break():
    doc.add_page_break()


# ============================================================
# Konten BAB III
# ============================================================

# Cover-ish header
h("BAB III\nANALISIS DAN PERANCANGAN SISTEM", level=1)
p(
    "Bab ini menjabarkan landasan teori yang menopang pengembangan Sistem Informasi YBY NET (SI YBY NET) "
    "serta perancangan rinci yang meliputi perancangan sistem (use case, activity, sequence, class diagram), "
    "perancangan basis data (ERD dan struktur tabel PostgreSQL), dan perancangan layar antarmuka."
)

# ---------- 3.1 Tinjauan singkat (opsional, tidak diminta) ----------

# ---------- 3.2 Tinjauan Pustaka ----------
h("3.2 Tinjauan Pustaka", level=2)
p(
    "Tinjauan pustaka memuat dasar teori, teknologi, dan konsep yang menjadi pijakan pembangunan SI YBY NET. "
    "Pemilihan setiap teknologi mengacu pada kebutuhan fungsional sistem (manajemen pelanggan ISP), kebutuhan "
    "non-fungsional (keandalan operasi terhadap perangkat MikroTik), dan praktik rekayasa perangkat lunak "
    "modern berbasis web."
)

h("3.2.1 Sistem Informasi dan ISP RT/RW Net", level=3)
p(
    "Sistem informasi (SI) adalah kombinasi perangkat keras, perangkat lunak, prosedur, dan sumber daya manusia "
    "yang dirancang untuk mengumpulkan, memproses, menyimpan, dan menyajikan data sehingga menghasilkan "
    "informasi yang berguna bagi pengambil keputusan. Pada konteks operator RT/RW Net, sistem informasi membantu "
    "mendigitalisasi proses bisnis manajemen pelanggan, penagihan, pengelolaan paket bandwidth, dan pencatatan "
    "tiket layanan yang sebelumnya banyak dilakukan secara manual atau melalui spreadsheet."
)
p(
    "RT/RW Net adalah model penyelenggara jasa internet skala lokal pada tingkat rukun tetangga atau rukun warga "
    "yang umumnya beroperasi sebagai reseller bandwidth dari Internet Service Provider (ISP) berlisensi. "
    "Karakteristik bisnisnya mencakup jumlah pelanggan kecil hingga menengah, rentang geografis terbatas, "
    "dan kebutuhan otomatisasi terhadap perangkat router (umumnya MikroTik) untuk pengaturan kuota dan kecepatan."
)

h("3.2.2 Next.js 14 dan React Server Components", level=3)
p(
    "Next.js adalah framework React yang menyediakan rendering hibrida (server-side rendering, static site "
    "generation, dan client-side rendering) di atas runtime Node.js. Pada versi 14 dengan App Router, Next.js "
    "memperkenalkan paradigma React Server Components (RSC) yang memungkinkan komponen dirender langsung di "
    "server sehingga ukuran bundle JavaScript di sisi klien menjadi lebih kecil dan performa awal halaman lebih "
    "cepat. SI YBY NET memanfaatkan struktur direktori berbasis App Router (src/app) dengan pemisahan grup rute "
    "antara halaman publik (group (auth)) dan halaman dashboard (group (dashboard))."
)

h("3.2.3 PostgreSQL dan Drizzle ORM", level=3)
p(
    "PostgreSQL adalah sistem manajemen basis data relasional bersifat open source dengan dukungan lengkap "
    "terhadap standar SQL, integritas referensial, transaksi ACID, dan tipe data lanjutan termasuk ENUM, JSONB, "
    "dan tipe waktu. Pemilihan PostgreSQL untuk SI YBY NET didasarkan pada kebutuhan integritas data finansial "
    "(pembayaran, laporan keuangan) dan kebutuhan tipe enumerasi yang ketat (status pelanggan, status tiket)."
)
p(
    "Drizzle ORM adalah Object-Relational Mapper (ORM) bertipe TypeScript-first yang memetakan tabel basis data "
    "ke definisi skema TypeScript, menghasilkan query type-safe, dan menyediakan migration toolkit "
    "(drizzle-kit) untuk menghasilkan dan menjalankan migrasi SQL secara deterministik."
)

h("3.2.4 NextAuth (Auth.js) dan Role-Based Access Control", level=3)
p(
    "NextAuth, kini dikenal sebagai Auth.js, adalah pustaka autentikasi untuk Next.js yang mendukung berbagai "
    "provider mulai dari OAuth, magic link, hingga Credentials. SI YBY NET menggunakan provider Credentials "
    "dengan strategi sesi JSON Web Token (JWT) yang menyimpan klaim id, email, name, dan role. Token JWT "
    "diverifikasi pada middleware (src/middleware.ts) sebelum permintaan mencapai halaman atau API route."
)
p(
    "Role-Based Access Control (RBAC) adalah pola pengamanan dengan memberi akses berdasarkan peran. SI YBY NET "
    "mendefinisikan tiga peran—admin, operator, dan technician—dengan pemetaan rute yang dipusatkan dalam "
    "objek roleAccess pada middleware sehingga setiap penambahan halaman cukup diatur dalam satu tempat."
)

h("3.2.5 MikroTik RouterOS API dan Manajemen PPPoE", level=3)
p(
    "MikroTik RouterOS adalah sistem operasi router yang menyediakan API biner pada port default 8728 (atau "
    "8729 untuk koneksi berenkripsi). API ini memungkinkan klien eksternal melakukan operasi seperti membuat "
    "pengguna PPPoE (/ppp/secret/add), menonaktifkan akun (/ppp/secret/set disabled=yes), memutus sesi aktif "
    "(/ppp/active/remove), serta mengatur batas bandwidth melalui simple queue (/queue/simple/add)."
)
p(
    "Pada SI YBY NET, integrasi dilakukan melalui pustaka Node.js node-routeros yang menerjemahkan perintah "
    "RouterOS API ke pemanggilan TypeScript. Modul integrasi dipecah menjadi client.ts (koneksi), pool.ts "
    "(connection pool dengan cache singkat), pppoe.ts (manajemen pengguna), queue.ts (manajemen bandwidth), "
    "dan monitor.ts (pemantauan koneksi aktif)."
)

h("3.2.6 BullMQ dan Redis untuk Pemrosesan Asinkron", level=3)
p(
    "BullMQ adalah pustaka antrian tugas (job queue) berbasis Redis untuk Node.js yang menyediakan jaminan "
    "pengiriman pesan, retry dengan exponential backoff, dan kontrol konkurensi. SI YBY NET menggunakan BullMQ "
    "untuk memisahkan operasi terhadap perangkat MikroTik dari siklus permintaan-respons HTTP. Hal ini penting "
    "karena perintah RouterOS dapat memerlukan beberapa detik atau gagal akibat kondisi jaringan—jika dieksekusi "
    "sinkron pada API route, latensi pengguna akan tinggi dan kegagalan akan diteruskan langsung ke pengguna "
    "akhir tanpa mekanisme retry."
)
p(
    "Dua antrian utama didefinisikan: mikrotik-queue (konkurensi 1, untuk menghindari race condition pada "
    "RouterOS) dan notification-queue (konkurensi 5, untuk pengiriman pesan WhatsApp paralel)."
)

h("3.2.7 Fonnte WhatsApp Gateway", level=3)
p(
    "Fonnte adalah penyedia gateway WhatsApp Business API tidak resmi yang banyak digunakan oleh UKM Indonesia. "
    "Komunikasi dilakukan melalui HTTP POST ke endpoint /send dengan header Authorization: Bearer <token>, "
    "menerima payload target (nomor telepon) dan message (teks pesan). SI YBY NET memanfaatkan Fonnte untuk "
    "mengirim notifikasi konfirmasi pembayaran dan pengingat tagihan dalam Bahasa Indonesia."
)

h("3.2.8 Unified Modeling Language (UML) dan Entity Relationship Diagram (ERD)", level=3)
p(
    "Unified Modeling Language (UML) adalah notasi standar untuk pemodelan perangkat lunak berorientasi objek. "
    "Empat diagram UML digunakan dalam laporan ini:"
)
bullet("Use Case Diagram menggambarkan interaksi antara aktor (pengguna sistem) dengan fungsionalitas yang disediakan sistem.")
bullet("Activity Diagram menggambarkan alur kerja (workflow) dari sebuah proses bisnis termasuk percabangan, paralelisme, dan penggabungan kembali.")
bullet("Sequence Diagram menggambarkan interaksi antar objek atau komponen dalam urutan waktu, sangat cocok untuk mendokumentasikan flow API/RPC.")
bullet("Class Diagram menggambarkan struktur statis sistem berupa kelas, atribut, operasi, dan relasi antar kelas.")
p(
    "Entity Relationship Diagram (ERD) adalah notasi grafis untuk memodelkan basis data relasional, yang "
    "menggambarkan entitas (tabel), atribut (kolom), dan relasi antar entitas dengan kardinalitas (one-to-one, "
    "one-to-many, many-to-many)."
)

h("3.2.9 Pustaka Pendukung Antarmuka", level=3)
p(
    "Tampilan antarmuka SI YBY NET dibangun dengan Tailwind CSS sebagai utility-first CSS framework dan "
    "shadcn/ui sebagai pustaka komponen siap pakai (button, dialog, table, dll.) yang dapat di-customize. "
    "TanStack Table digunakan untuk fitur tabel data lanjutan (sorting, filtering, pagination), Recharts untuk "
    "grafik laporan keuangan dan bandwidth, React Hook Form untuk pengelolaan state form, dan Zod untuk skema "
    "validasi yang dapat dipakai bersama di sisi klien dan server."
)

page_break()

# ---------- 3.3 Perancangan Sistem ----------
h("3.3 Perancangan Sistem", level=2)
p(
    "Perancangan sistem mencakup pemodelan kebutuhan fungsional menggunakan empat diagram UML: use case "
    "diagram, activity diagram, sequence diagram, dan class diagram. Pemodelan ini disusun berdasarkan analisis "
    "alur bisnis YBY NET dan kebutuhan teknis integrasi MikroTik serta gateway WhatsApp."
)

# 3.3.1 Use Case
h("3.3.1 Use Case Diagram", level=3)
p(
    "Use case diagram pada Gambar 3.1 menggambarkan interaksi empat aktor dengan SI YBY NET. Aktor admin "
    "memiliki akses penuh terhadap semua fungsi termasuk konfigurasi sistem dan manajemen pengguna. Operator "
    "berfokus pada operasi harian seperti pencatatan pelanggan dan pembayaran. Teknisi memiliki akses terbatas "
    "ke modul monitoring bandwidth, tiket layanan, dan activity log. Pelanggan/calon pelanggan dapat mengajukan "
    "permintaan layanan melalui formulir publik tanpa otentikasi."
)
add_image(DIAG / "usecase.png", width_in=6.3, cap="Gambar 3.1 Use Case Diagram SI YBY NET")

p("Tabel 3.1 menyajikan deskripsi rinci enam use case utama.", indent_first=False)
make_table(
    headers=["Nama Use Case", "Aktor", "Prakondisi", "Alur Normal", "Alur Alternatif"],
    rows=[
        [
            "Login",
            "Admin, Operator, Teknisi",
            "Pengguna telah terdaftar di tabel users.",
            "(1) Pengguna membuka /login. (2) Sistem menampilkan form. (3) Pengguna memasukkan email & password. (4) Sistem memverifikasi melalui NextAuth Credentials. (5) Sistem mengeluarkan JWT dan mengarahkan ke /.",
            "Bila kredensial salah, sistem menampilkan pesan error dan tetap di halaman /login.",
        ],
        [
            "Kelola Pelanggan",
            "Admin, Operator",
            "Pengguna telah login dengan role admin/operator.",
            "(1) Pengguna membuka /customers. (2) Sistem menampilkan daftar. (3) Pengguna memilih aksi tambah/edit. (4) Mengisi form dan submit. (5) Sistem memvalidasi dan menyimpan ke DB. (6) Bila ada pppoe_username, sistem mengantrekan job CREATE_PPPOE/SET_QUEUE.",
            "Validasi gagal → kembali ke form dengan pesan kesalahan per field. Job MikroTik gagal setelah 3 retry → tercatat di log dan ditandai gagal.",
        ],
        [
            "Catat Pembayaran",
            "Admin, Operator",
            "Pelanggan terdaftar dan aktif.",
            "(1) Pengguna membuka /payments/new. (2) Memilih pelanggan dan periode. (3) Mengisi nominal, tanggal, metode. (4) Submit. (5) Sistem membuat invoice, menyimpan record, mengantrekan WhatsAppBilling. (6) Menampilkan kuitansi digital.",
            "Pelanggan belum aktif → tampilkan peringatan. Pengiriman WhatsApp gagal → job retry; pembayaran tetap tercatat.",
        ],
        [
            "Kelola Tiket Layanan",
            "Admin, Operator, Teknisi",
            "Tiket telah diajukan oleh pelanggan.",
            "(1) Pengguna membuka /service-requests. (2) Memilih tiket. (3) Mengubah status (open → in_progress → resolved). (4) Mengisi catatan admin/teknisi. (5) Sistem mencatat resolvedAt.",
            "Tiket berjenis 'unsubscribe' yang di-approve memicu job DELETE_PPPOE secara otomatis.",
        ],
        [
            "Monitor Bandwidth",
            "Admin, Operator, Teknisi",
            "Koneksi MikroTik aktif.",
            "(1) Pengguna membuka /bandwidth. (2) Sistem memanggil /api/bandwidth. (3) API mengambil data dari MikroTik pool (cache 10 detik). (4) Menampilkan tabel sesi aktif, hitungan online/offline, dan grafik throughput.",
            "Koneksi MikroTik gagal → API mengembalikan 503; UI menampilkan pesan 'Tidak dapat menghubungi router'.",
        ],
        [
            "Submit Pengajuan Layanan",
            "Pelanggan / calon pelanggan (publik)",
            "—",
            "(1) Pengunjung membuka halaman publik service request. (2) Memilih jenis pengajuan. (3) Mengisi nama, telepon, deskripsi. (4) Submit. (5) Sistem menyimpan tiket dengan status 'open'.",
            "Validasi gagal → form menampilkan kesalahan. Tiket pending menumpuk → admin menerima daftar pada dashboard.",
        ],
    ],
    col_widths_cm=[3, 2.2, 2.2, 5, 3.5],
)

# 3.3.2 Activity
h("3.3.2 Activity Diagram", level=3)
p(
    "Empat activity diagram disusun untuk merepresentasikan alur kerja inti SI YBY NET: pendaftaran pelanggan, "
    "pencatatan pembayaran, suspend bandwidth (perubahan status), dan pengajuan/penanganan tiket layanan."
)

h("a. Activity Pendaftaran Pelanggan", level=4)
p(
    "Gambar 3.2 menampilkan alur pendaftaran pelanggan baru. Diagram ini menunjukkan transisi dari aksi sinkron "
    "(simpan ke PostgreSQL) ke aksi asinkron (provisioning PPPoE di MikroTik) yang ditangani oleh worker BullMQ."
)
add_image(DIAG / "activity_register_customer.png", width_in=6.0, cap="Gambar 3.2 Activity Diagram Pendaftaran Pelanggan")

h("b. Activity Pencatatan Pembayaran", level=4)
p(
    "Gambar 3.3 menggambarkan proses pencatatan pembayaran beserta percabangan ke notification worker untuk "
    "pengiriman konfirmasi WhatsApp."
)
add_image(DIAG / "activity_payment.png", width_in=6.0, cap="Gambar 3.3 Activity Diagram Pencatatan Pembayaran")

h("c. Activity Suspend / Aktifkan Pelanggan", level=4)
p(
    "Gambar 3.4 menggambarkan alur ketika admin/operator mengubah status pelanggan, yang otomatis memicu "
    "perintah ke MikroTik untuk men-disable PPPoE secret dan memutus sesi aktif."
)
add_image(DIAG / "activity_suspend_bandwidth.png", width_in=6.0, cap="Gambar 3.4 Activity Diagram Suspend Pelanggan & Bandwidth")

h("d. Activity Pengajuan dan Penanganan Tiket Layanan", level=4)
p(
    "Gambar 3.5 menggambarkan alur tiket layanan dari sisi pelanggan (pengajuan) hingga sisi admin/teknisi "
    "(diagnosa, eskalasi, dan penyelesaian)."
)
add_image(DIAG / "activity_service_request.png", width_in=6.0, cap="Gambar 3.5 Activity Diagram Pengajuan & Penanganan Tiket")

# 3.3.3 Sequence
h("3.3.3 Sequence Diagram", level=3)
p(
    "Sequence diagram menggambarkan interaksi antar komponen perangkat lunak (browser, API route, validator, "
    "DB, antrian, worker, dan layanan eksternal) dalam urutan waktu. Empat sequence diagram dibuat untuk skenario "
    "inti."
)

h("a. Sequence Login", level=4)
add_image(DIAG / "sequence_login.png", width_in=6.0, cap="Gambar 3.6 Sequence Diagram Login")

h("b. Sequence Pendaftaran Pelanggan dengan Provisioning PPPoE", level=4)
add_image(DIAG / "sequence_register_customer.png", width_in=6.0, cap="Gambar 3.7 Sequence Diagram Pendaftaran Pelanggan")

h("c. Sequence Pencatatan Pembayaran dan Notifikasi WhatsApp", level=4)
add_image(DIAG / "sequence_payment_notification.png", width_in=6.0, cap="Gambar 3.8 Sequence Diagram Pembayaran & Notifikasi")

h("d. Sequence Monitor Bandwidth Real-time", level=4)
add_image(DIAG / "sequence_bandwidth_monitor.png", width_in=6.0, cap="Gambar 3.9 Sequence Diagram Monitor Bandwidth")

h("e. Sequence Suspend Pelanggan", level=4)
add_image(DIAG / "sequence_suspend_pppoe.png", width_in=6.0, cap="Gambar 3.10 Sequence Diagram Suspend PPPoE")

# 3.3.4 Class Diagram
h("3.3.4 Class Diagram", level=3)
p(
    "Class diagram pada Gambar 3.11 memodelkan entitas domain SI YBY NET beserta relasinya. Selain entitas "
    "data (User, Customer, Package, Payment, ServiceRequest, ActivityLog, Setting), terdapat tiga kelas helper "
    "yang merepresentasikan integrasi eksternal: MikrotikClient, BullQueueProducer, dan WhatsAppService."
)
add_image(DIAG / "class_diagram.png", width_in=6.3, cap="Gambar 3.11 Class Diagram SI YBY NET")

page_break()

# ---------- 3.4 Perancangan Database ----------
h("3.4 Perancangan Basis Data", level=2)
p(
    "SI YBY NET menggunakan PostgreSQL 16 sebagai sistem basis data. Skema dirancang dengan prinsip "
    "normalisasi 3NF, integritas referensial via foreign key, dan tipe ENUM untuk field berstatus. "
    "Bagian ini memuat Entity Relationship Diagram (ERD), spesifikasi tabel, dan Data Definition Language "
    "(DDL) lengkap."
)

h("3.4.1 Entity Relationship Diagram", level=3)
p(
    "Gambar 3.12 menyajikan ERD SI YBY NET. Terdapat tujuh entitas: users (pengguna sistem), packages "
    "(paket layanan), customers (pelanggan), payments (pembayaran), service_requests (tiket layanan), "
    "activity_logs (log aktivitas pelanggan dari MikroTik), dan settings (key-value konfigurasi sistem)."
)
add_image(DIAG / "erd.png", width_in=6.3, cap="Gambar 3.12 Entity Relationship Diagram SI YBY NET")

p("Relasi antar entitas dirangkum dalam Tabel 3.2.", indent_first=False)
make_table(
    headers=["Entitas Induk", "Entitas Anak", "Kolom FK", "Kardinalitas", "Keterangan"],
    rows=[
        ["packages", "customers", "package_id", "1 — N", "Satu paket dapat dimiliki banyak pelanggan."],
        ["customers", "payments", "customer_id", "1 — N", "Satu pelanggan dapat memiliki banyak pembayaran."],
        ["customers", "service_requests", "customer_id", "1 — N (opsional)", "Tiket dapat berasal dari pelanggan terdaftar atau tamu (NULL)."],
        ["customers", "activity_logs", "customer_id", "1 — N", "Setiap log login/logout terkait satu pelanggan."],
        ["users", "payments", "received_by", "1 — N", "Mencatat operator yang menerima pembayaran."],
        ["users", "service_requests", "handled_by", "1 — N (opsional)", "Mencatat staf yang menangani tiket."],
    ],
    col_widths_cm=[3, 3, 2.5, 3, 4.5],
)

# 3.4.2 Spesifikasi Tabel
h("3.4.2 Spesifikasi Tabel", level=3)
p(
    "Bagian ini memuat spesifikasi setiap tabel dalam basis data SI YBY NET. Tipe data mengikuti dialek "
    "PostgreSQL; varchar(N) berarti karakter variabel maksimal N, numeric(p,s) berarti angka presisi p "
    "dengan s digit desimal."
)

table_specs = [
    (
        "Tabel 3.3 Spesifikasi Tabel users",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik pengguna, auto-increment."],
            ["name", "varchar(255)", "NOT NULL", "Nama lengkap pengguna."],
            ["email", "varchar(255)", "NOT NULL, UNIQUE", "Email login (kunci unik)."],
            ["password", "varchar(255)", "NOT NULL", "Hash bcrypt dari password."],
            ["role", "user_role", "NOT NULL, DEFAULT 'operator'", "Peran: admin/operator/technician."],
            ["created_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp pembuatan akun."],
        ],
    ),
    (
        "Tabel 3.4 Spesifikasi Tabel packages",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik paket."],
            ["name", "varchar(100)", "NOT NULL", "Nama paket (mis. 'Paket 20 Mbps')."],
            ["speed", "integer", "NOT NULL", "Kecepatan dalam satuan Mbps."],
            ["price", "numeric(12,2)", "NOT NULL", "Harga per bulan dalam Rupiah."],
            ["queue_target", "varchar(50)", "NULL", "Target Simple Queue MikroTik (mis. '20M/20M')."],
            ["is_active", "boolean", "NOT NULL, DEFAULT true", "Apakah paket masih ditawarkan."],
            ["created_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp pembuatan."],
        ],
    ),
    (
        "Tabel 3.5 Spesifikasi Tabel customers",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik pelanggan."],
            ["customer_id", "varchar(20)", "NOT NULL, UNIQUE", "ID pelanggan format YBY-YYMM-NNNN."],
            ["name", "varchar(255)", "NOT NULL", "Nama lengkap pelanggan."],
            ["address", "text", "NOT NULL", "Alamat lengkap."],
            ["phone", "varchar(20)", "NOT NULL", "Nomor telepon (08xxxxxxxxxx)."],
            ["email", "varchar(255)", "NULL", "Email pelanggan (opsional)."],
            ["nik", "varchar(16)", "NULL", "Nomor Induk Kependudukan."],
            ["package_id", "integer", "NOT NULL, FK packages(id)", "Paket yang dilanggan."],
            ["registration_date", "date", "NOT NULL", "Tanggal pendaftaran."],
            ["active_until", "date", "NOT NULL", "Tanggal kadaluarsa langganan."],
            ["status", "customer_status", "NOT NULL, DEFAULT 'active'", "Status: active/inactive/suspended/terminated."],
            ["pppoe_username", "varchar(100)", "NULL", "Username PPPoE pada MikroTik."],
            ["onu_sn", "varchar(50)", "NULL", "Serial number ONU (untuk fiber)."],
            ["notes", "text", "NULL", "Catatan internal."],
            ["created_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp pembuatan."],
            ["updated_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp perubahan terakhir."],
        ],
    ),
    (
        "Tabel 3.6 Spesifikasi Tabel payments",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik pembayaran."],
            ["invoice_number", "varchar(30)", "NOT NULL, UNIQUE", "Nomor invoice format INV-YYMM-NNNNN."],
            ["customer_id", "integer", "NOT NULL, FK customers(id)", "Pelanggan yang membayar."],
            ["amount", "numeric(12,2)", "NOT NULL", "Nominal pembayaran."],
            ["payment_date", "date", "NOT NULL", "Tanggal pembayaran."],
            ["payment_method", "varchar(20)", "NOT NULL", "Metode: cash/transfer/qris."],
            ["period_month", "varchar(7)", "NOT NULL", "Periode tagihan format YYYY-MM."],
            ["status", "payment_status", "NOT NULL, DEFAULT 'pending'", "Status: paid/pending/overdue/cancelled."],
            ["notes", "text", "NULL", "Catatan tambahan."],
            ["received_by", "integer", "NOT NULL, FK users(id)", "Staf yang menerima pembayaran."],
            ["created_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp pembuatan."],
        ],
    ),
    (
        "Tabel 3.7 Spesifikasi Tabel service_requests",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik tiket."],
            ["ticket_number", "varchar(30)", "NOT NULL, UNIQUE", "Nomor tiket format TKT-YYMM-NNNN."],
            ["type", "request_type", "NOT NULL", "Jenis: new_installation/upgrade_downgrade/trouble_ticket/relocation/unsubscribe."],
            ["customer_id", "integer", "NULL, FK customers(id)", "Pelanggan terkait (NULL untuk tamu)."],
            ["name", "varchar(255)", "NOT NULL", "Nama pemohon."],
            ["phone", "varchar(20)", "NOT NULL", "Telepon pemohon."],
            ["description", "text", "NOT NULL", "Deskripsi keluhan/permintaan."],
            ["status", "request_status", "NOT NULL, DEFAULT 'open'", "Status tiket."],
            ["admin_notes", "text", "NULL", "Catatan admin/teknisi."],
            ["handled_by", "integer", "NULL, FK users(id)", "Staf penanggung jawab."],
            ["resolved_at", "timestamp", "NULL", "Timestamp penyelesaian."],
            ["created_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp pengajuan."],
            ["updated_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp perubahan terakhir."],
        ],
    ),
    (
        "Tabel 3.8 Spesifikasi Tabel activity_logs",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik log."],
            ["customer_id", "integer", "NOT NULL, FK customers(id)", "Pelanggan terkait."],
            ["action", "varchar(50)", "NOT NULL", "Jenis aksi: login/logout/dll."],
            ["duration", "integer", "NULL", "Durasi sesi dalam detik."],
            ["bytes_in", "numeric(15,0)", "NULL", "Total byte download."],
            ["bytes_out", "numeric(15,0)", "NULL", "Total byte upload."],
            ["timestamp", "timestamp", "NOT NULL, DEFAULT now()", "Waktu kejadian."],
        ],
    ),
    (
        "Tabel 3.9 Spesifikasi Tabel settings",
        [
            ["id", "serial", "PRIMARY KEY", "ID unik."],
            ["key", "varchar(100)", "NOT NULL, UNIQUE", "Kunci konfigurasi (mis. 'mikrotik_host')."],
            ["value", "varchar(1000)", "NOT NULL, DEFAULT ''", "Nilai konfigurasi."],
            ["updated_at", "timestamp", "NOT NULL, DEFAULT now()", "Timestamp perubahan."],
        ],
    ),
]

for title, rows in table_specs:
    p(title, bold=True, indent_first=False, align="left")
    make_table(
        headers=["Kolom", "Tipe Data", "Constraint", "Keterangan"],
        rows=rows,
        col_widths_cm=[3.5, 3, 4, 4],
    )

# 3.4.3 DDL
h("3.4.3 Data Definition Language (DDL) PostgreSQL", level=3)
p(
    "Berikut adalah skrip DDL lengkap untuk membuat seluruh tipe enum dan tabel pada basis data SI YBY NET. "
    "Skrip ini diturunkan dari migration file yang dihasilkan oleh drizzle-kit dan telah diterapkan pada "
    "lingkungan produksi."
)
code_block(
    """-- ENUM TYPES
CREATE TYPE user_role AS ENUM ('admin', 'operator', 'technician');
CREATE TYPE customer_status AS ENUM ('active', 'inactive', 'suspended', 'terminated');
CREATE TYPE payment_status AS ENUM ('paid', 'pending', 'overdue', 'cancelled');
CREATE TYPE request_type AS ENUM (
    'new_installation', 'upgrade_downgrade', 'trouble_ticket',
    'relocation', 'unsubscribe'
);
CREATE TYPE request_status AS ENUM (
    'open', 'in_progress', 'approved', 'rejected', 'resolved', 'closed'
);

-- TABLES
CREATE TABLE users (
    id          serial PRIMARY KEY,
    name        varchar(255) NOT NULL,
    email       varchar(255) NOT NULL UNIQUE,
    password    varchar(255) NOT NULL,
    role        user_role NOT NULL DEFAULT 'operator',
    created_at  timestamp NOT NULL DEFAULT now()
);

CREATE TABLE packages (
    id            serial PRIMARY KEY,
    name          varchar(100) NOT NULL,
    speed         integer NOT NULL,
    price         numeric(12,2) NOT NULL,
    queue_target  varchar(50),
    is_active     boolean NOT NULL DEFAULT true,
    created_at    timestamp NOT NULL DEFAULT now()
);

CREATE TABLE customers (
    id                serial PRIMARY KEY,
    customer_id       varchar(20) NOT NULL UNIQUE,
    name              varchar(255) NOT NULL,
    address           text NOT NULL,
    phone             varchar(20) NOT NULL,
    email             varchar(255),
    nik               varchar(16),
    package_id        integer NOT NULL REFERENCES packages(id),
    registration_date date NOT NULL,
    active_until      date NOT NULL,
    status            customer_status NOT NULL DEFAULT 'active',
    pppoe_username    varchar(100),
    onu_sn            varchar(50),
    notes             text,
    created_at        timestamp NOT NULL DEFAULT now(),
    updated_at        timestamp NOT NULL DEFAULT now()
);

CREATE TABLE payments (
    id              serial PRIMARY KEY,
    invoice_number  varchar(30) NOT NULL UNIQUE,
    customer_id     integer NOT NULL REFERENCES customers(id),
    amount          numeric(12,2) NOT NULL,
    payment_date    date NOT NULL,
    payment_method  varchar(20) NOT NULL,
    period_month    varchar(7) NOT NULL,
    status          payment_status NOT NULL DEFAULT 'pending',
    notes           text,
    received_by     integer NOT NULL REFERENCES users(id),
    created_at      timestamp NOT NULL DEFAULT now()
);

CREATE TABLE service_requests (
    id             serial PRIMARY KEY,
    ticket_number  varchar(30) NOT NULL UNIQUE,
    type           request_type NOT NULL,
    customer_id    integer REFERENCES customers(id),
    name           varchar(255) NOT NULL,
    phone          varchar(20) NOT NULL,
    description    text NOT NULL,
    status         request_status NOT NULL DEFAULT 'open',
    admin_notes    text,
    handled_by     integer REFERENCES users(id),
    resolved_at    timestamp,
    created_at     timestamp NOT NULL DEFAULT now(),
    updated_at     timestamp NOT NULL DEFAULT now()
);

CREATE TABLE activity_logs (
    id           serial PRIMARY KEY,
    customer_id  integer NOT NULL REFERENCES customers(id),
    action       varchar(50) NOT NULL,
    duration     integer,
    bytes_in     numeric(15,0),
    bytes_out    numeric(15,0),
    timestamp    timestamp NOT NULL DEFAULT now()
);

CREATE TABLE settings (
    id          serial PRIMARY KEY,
    key         varchar(100) NOT NULL UNIQUE,
    value       varchar(1000) NOT NULL DEFAULT '',
    updated_at  timestamp NOT NULL DEFAULT now()
);

-- INDEX TAMBAHAN UNTUK PERFORMA QUERY
CREATE INDEX idx_customers_status ON customers(status);
CREATE INDEX idx_customers_pppoe ON customers(pppoe_username);
CREATE INDEX idx_payments_customer ON payments(customer_id);
CREATE INDEX idx_payments_period ON payments(period_month);
CREATE INDEX idx_service_status ON service_requests(status);
CREATE INDEX idx_activity_customer_ts ON activity_logs(customer_id, timestamp DESC);
"""
)

page_break()

# ---------- 3.5 Perancangan Layar ----------
h("3.5 Perancangan Layar (User Interface)", level=2)
p(
    "Perancangan layar disusun sebagai panduan tata letak antarmuka SI YBY NET. Setiap layar dideskripsikan "
    "berdasarkan tujuan, role yang berhak mengakses, struktur tata letak (header/sidebar/konten utama/aksi), "
    "wireframe ASCII sebagai sketsa kasar, dan daftar field beserta validasi yang merujuk pada skema Zod. "
    "Wireframe ini dapat dikembangkan menjadi mockup hi-fi menggunakan Figma atau Canva."
)
p(
    "Seluruh layar mengikuti tata letak konsisten: di sebelah kiri terdapat sidebar navigasi (kecuali halaman "
    "/login), bagian atas terdapat header dengan informasi pengguna dan tombol logout, dan area utama berisi "
    "konten halaman. Tema warna menggunakan palet biru-putih dengan aksen hijau untuk status aktif dan merah "
    "untuk status terancam (suspended/overdue)."
)

screens = [
    {
        "title": "3.5.1 Halaman Login",
        "url": "/login",
        "role": "Publik (semua orang dapat mengakses)",
        "purpose": "Autentikasi pengguna ke dalam sistem.",
        "structure": [
            "Background gambar (public/bg-ocean.jpg) dengan overlay gelap.",
            "Card berukuran ±400px berisi logo YBY NET, judul 'Masuk', form email & password, tombol 'Masuk'.",
            "Footer kecil dengan copyright.",
        ],
        "ascii": """
+------------------------------------------------------------+
|              [bg-ocean.jpg dengan overlay]                 |
|                                                            |
|              +----------------------------+                |
|              |       [Logo YBY NET]       |                |
|              |          MASUK             |                |
|              | -------------------------- |                |
|              | Email                      |                |
|              | [____________________]     |                |
|              | Password                   |                |
|              | [____________________]     |                |
|              |                            |                |
|              |    [        MASUK        ] |                |
|              +----------------------------+                |
|                                                            |
|                  (c) 2026 YBY NET                          |
+------------------------------------------------------------+""",
        "fields": [
            ["email", "email", "Wajib, format email valid", "loginSchema"],
            ["password", "password", "Wajib, minimal 6 karakter", "loginSchema"],
        ],
    },
    {
        "title": "3.5.2 Halaman Dashboard (Beranda)",
        "url": "/",
        "role": "admin, operator, technician",
        "purpose": "Menampilkan ringkasan operasional ISP secara real-time.",
        "structure": [
            "Sidebar navigasi (kiri) dengan menu: Beranda, Pelanggan, Paket, Pembayaran, Tiket, Bandwidth, Activity Log, Laporan, Pengaturan.",
            "Header (atas) dengan nama pengguna, role badge, dan tombol logout.",
            "Grid kartu statistik 4 kolom: Total Pelanggan, Pelanggan Aktif, Tunggakan Bulan Ini, Tiket Terbuka.",
            "Dua chart sejajar: tren pendapatan 12 bulan (LineChart Recharts) dan distribusi status pelanggan (PieChart).",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Header: User: Admin (admin)         [Logout]      |
| [Logo] +---------------------------------------------------+
|========|                                                   |
| Beranda|  +---------+ +---------+ +---------+ +---------+ |
| Plgn.  |  | TOTAL   | | AKTIF   | | TUNGGAK | | TIKET   | |
| Paket  |  |  148    | |  131    | |   12    | |    4    | |
| Bayar  |  +---------+ +---------+ +---------+ +---------+ |
| Tiket  |                                                   |
| BW     |  +-----------------------+  +-----------------+   |
| Log    |  | Pendapatan 12 Bulan   |  | Status Plgn.    |   |
| Lapor. |  |  (Line Chart)         |  | (Pie Chart)     |   |
| Setup  |  |                       |  |                 |   |
|        |  +-----------------------+  +-----------------+   |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.3 Halaman Daftar Pelanggan",
        "url": "/customers",
        "role": "admin, operator",
        "purpose": "Menampilkan daftar pelanggan dengan fitur filter, sort, paging, dan aksi massal.",
        "structure": [
            "Bagian atas: judul 'Pelanggan', tombol 'Tambah Pelanggan' (kanan atas), tombol 'Export CSV'.",
            "Filter bar: search by name/customer_id/phone, filter status, filter paket.",
            "Tabel data (TanStack Table) kolom: ID, Nama, Telepon, Paket, Status, Aktif Hingga, Aksi.",
            "Footer tabel: pagination (10/25/50 per halaman), tombol bulk action (suspend/aktifkan).",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Pelanggan              [+ Tambah]  [Export CSV]   |
|        +---------------------------------------------------+
|        | [Cari...] [Status v] [Paket v]                    |
| (sb)   +---------------------------------------------------+
|        | [ ] ID         Nama         Tlp   Paket  Status   |
|        | [ ] YBY-2603-1 Pak Budi    08xx 20Mbps Aktif [..] |
|        | [ ] YBY-2603-2 Bu Ani      08xx 10Mbps Suspd [..] |
|        | [ ] ...                                           |
|        +---------------------------------------------------+
|        | [Bulk: Suspend|Aktifkan]   << 1 2 3 ... 15 >>     |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.4 Halaman Tambah / Edit Pelanggan",
        "url": "/customers/new, /customers/[id]/edit",
        "role": "admin, operator",
        "purpose": "Form input data pelanggan baru atau edit data eksisting.",
        "structure": [
            "Form dua kolom dengan section: Identitas, Kontak, Layanan.",
            "Section Identitas: Nama, NIK (opsional), Alamat.",
            "Section Kontak: Telepon, Email (opsional).",
            "Section Layanan: Paket (Select), Tanggal Daftar, Aktif Hingga, Status (hanya pada edit), PPPoE Username, ONU SN, Catatan.",
            "Action button: Simpan (primary), Batal.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | < Kembali     Tambah Pelanggan                    |
|        +---------------------------------------------------+
|        | == Identitas ==           == Kontak ==            |
|        | Nama* [______________]    Telp*  [____________]   |
| (sb)   | NIK   [______________]    Email  [____________]   |
|        | Alamat                                            |
|        | [_____________________________________________]   |
|        | == Layanan ==                                     |
|        | Paket* [Select v]      Daftar* [date]             |
|        | PPPoE  [______________]  ONU SN [____________]    |
|        | Catatan                                           |
|        | [_____________________________________________]   |
|        |                                                   |
|        |               [Batal]   [    SIMPAN    ]          |
+--------+---------------------------------------------------+""",
        "fields": [
            ["name", "text", "Wajib, minimal 2 karakter", "createCustomerSchema"],
            ["address", "textarea", "Wajib, minimal 5 karakter", "createCustomerSchema"],
            ["phone", "text", "Wajib, 10–15 digit format 08xxxxxxxxxx", "createCustomerSchema"],
            ["email", "email", "Opsional, format email valid", "createCustomerSchema"],
            ["nik", "text", "Opsional, tepat 16 digit numerik", "createCustomerSchema"],
            ["packageId", "select", "Wajib, integer positif", "createCustomerSchema"],
            ["pppoeUsername", "text", "Opsional, alfanumerik + dash/underscore", "createCustomerSchema"],
            ["onuSn", "text", "Opsional", "createCustomerSchema"],
            ["status", "select", "Wajib (saat edit): active/inactive/suspended/terminated", "updateCustomerSchema"],
            ["notes", "textarea", "Opsional", "createCustomerSchema"],
        ],
    },
    {
        "title": "3.5.5 Halaman Detail Pelanggan",
        "url": "/customers/[id]",
        "role": "admin, operator",
        "purpose": "Menampilkan detail lengkap pelanggan beserta riwayat pembayaran dan tiket.",
        "structure": [
            "Header: nama pelanggan, badge status, tombol Edit dan Suspend/Aktifkan.",
            "Card Identitas (kiri): nama, customer_id, NIK, alamat, telepon, email.",
            "Card Layanan (kanan): paket, kecepatan, harga, tanggal daftar, aktif hingga, PPPoE, ONU SN.",
            "Tab: Riwayat Pembayaran (tabel), Riwayat Tiket (tabel), Activity Log (tabel).",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | < Kembali  Pak Budi  [Aktif]   [Edit] [Suspend]   |
|        +---------------------------------------------------+
|        | +--- Identitas ---+ +--- Layanan ----+            |
|        | | YBY-2603-0001    | Paket 20 Mbps   |            |
| (sb)   | | NIK 32xxxxxx     | Rp 200.000      |            |
|        | | Alamat ...       | Aktif s/d 30/05 |            |
|        | | Tlp 08xxx        | PPPoE: budi-net |            |
|        | +------------------+ +-----------------+           |
|        | [ Pembayaran ] [ Tiket ] [ Activity Log ]         |
|        | +-----------------------------------------------+ |
|        | | INV-2603-001 | 2026-04-30 | Rp200.000 | Lunas  | |
|        | | INV-2603-002 | 2026-03-30 | Rp200.000 | Lunas  | |
|        | +-----------------------------------------------+ |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.6 Halaman Daftar Paket",
        "url": "/packages",
        "role": "admin, operator (form 'tambah paket' hanya admin)",
        "purpose": "Mengelola paket layanan yang ditawarkan.",
        "structure": [
            "Header: 'Paket Layanan', tombol 'Tambah Paket' (admin).",
            "Tabel: Nama, Kecepatan (Mbps), Harga, Queue Target, Pelanggan Berlangganan (count), Status, Aksi.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Paket Layanan                  [+ Tambah Paket]   |
|        +---------------------------------------------------+
|        | Nama        Speed  Harga      Queue    Plgn. Aksi |
| (sb)   | 10 Mbps     10     Rp150.000  10M/10M  42   [..]  |
|        | 20 Mbps     20     Rp250.000  20M/20M  68   [..]  |
|        | 50 Mbps     50     Rp500.000  50M/50M  21   [..]  |
+--------+---------------------------------------------------+""",
        "fields": [
            ["name", "text", "Wajib, min 2 karakter", "createPackageSchema"],
            ["speed", "number", "Wajib, integer positif (Mbps)", "createPackageSchema"],
            ["price", "number", "Wajib, non-negatif", "createPackageSchema"],
            ["queueTarget", "text", "Opsional, format \\d+[KMG]/\\d+[KMG] (mis. '20M/20M')", "createPackageSchema"],
            ["isActive", "checkbox", "Boolean, default true", "createPackageSchema"],
        ],
    },
    {
        "title": "3.5.7 Halaman Pembayaran dan Form Pembayaran",
        "url": "/payments, /payments/new",
        "role": "admin, operator",
        "purpose": "Mencatat pembayaran tagihan bulanan dan mengirim notifikasi WhatsApp.",
        "structure": [
            "Daftar: filter periode bulan, search pelanggan, tombol 'Catat Pembayaran'.",
            "Tabel: No. Invoice, Pelanggan, Periode, Nominal, Metode, Status, Tanggal, Aksi.",
            "Form: Pelanggan (combobox dengan auto-suggest), Periode (month picker), Nominal (auto-isi dari paket), Tanggal, Metode, Catatan.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Pembayaran        [Periode: 2026-05] [+ Catat]    |
|        +---------------------------------------------------+
|        | Invoice         Plgn        Periode   Nominal St. |
| (sb)   | INV-2605-00001  Pak Budi   2026-05  Rp200rb Lunas |
|        | INV-2605-00002  Bu Ani     2026-05  Rp150rb Lunas |
|        +---------------------------------------------------+
|
| Form Catat Pembayaran:
|   Pelanggan*   [Cari pelanggan v]
|   Periode*     [2026-05]
|   Nominal*     [Rp 200.000]
|   Tanggal*     [2026-05-08]
|   Metode*      ( ) Cash  (o) Transfer  ( ) QRIS
|   Catatan      [____________________]
|   [Batal]   [SIMPAN & KIRIM WHATSAPP]""",
        "fields": [
            ["customerId", "select", "Wajib, integer positif", "createPaymentSchema"],
            ["amount", "number", "Wajib, positif", "createPaymentSchema"],
            ["paymentDate", "date", "Wajib", "createPaymentSchema"],
            ["paymentMethod", "radio", "Wajib: cash/transfer/qris", "createPaymentSchema"],
            ["periodMonth", "month", "Wajib, format YYYY-MM", "createPaymentSchema"],
            ["notes", "textarea", "Opsional", "createPaymentSchema"],
        ],
    },
    {
        "title": "3.5.8 Halaman Tiket Layanan",
        "url": "/service-requests, /service-requests/[id]",
        "role": "admin, operator, technician",
        "purpose": "Mengelola tiket pengaduan, pemasangan baru, upgrade/downgrade, relokasi, dan unsubscribe.",
        "structure": [
            "Daftar (kiri-list, detail-kanan layout): kolom kiri menampilkan list tiket terbuka & in_progress dengan filter status & jenis, kolom kanan menampilkan detail tiket terpilih.",
            "Detail: nomor tiket, jenis, pemohon, deskripsi, riwayat status, form catatan admin/teknisi, tombol transisi status (Open → In Progress → Resolved → Closed).",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Tiket Layanan         [Filter: open v] [+ Baru]   |
|        +-----------+---------------------------------------+
|        | TKT-2605-1| TKT-2605-0001  Trouble Ticket         |
|        | Pak Budi  | Pemohon: Pak Budi (08xxxxxxxxxx)      |
|        | Trouble   | Diajukan: 2026-05-07 14:20            |
| (sb)   |           | Status: [Open v]   Penanggung: [- v]  |
|        | TKT-2605-2| Deskripsi:                            |
|        | Bu Ani    | Internet di rumah putus sejak pagi... |
|        | Pasang    | -- Catatan Admin/Teknisi --           |
|        |           | [_____________________________]       |
|        | TKT-2605-3| [        SIMPAN PERUBAHAN          ]  |
|        | Pak Joni  |                                       |
|        | Upgrade   |                                       |
+--------+-----------+---------------------------------------+""",
        "fields": [
            ["type", "select", "Wajib: new_installation/upgrade_downgrade/trouble_ticket/relocation/unsubscribe", "createServiceRequestSchema"],
            ["customerId", "select", "Opsional (kosong untuk tamu)", "createServiceRequestSchema"],
            ["name", "text", "Wajib, minimal 2 karakter", "createServiceRequestSchema"],
            ["phone", "text", "Wajib, ≥10 digit format 08xxxxxxxxxx", "createServiceRequestSchema"],
            ["description", "textarea", "Wajib, minimal 10 karakter", "createServiceRequestSchema"],
            ["status", "select (admin/operator)", "Wajib pada update", "updateServiceRequestSchema"],
            ["adminNotes", "textarea", "Opsional", "updateServiceRequestSchema"],
        ],
    },
    {
        "title": "3.5.9 Halaman Monitor Bandwidth",
        "url": "/bandwidth",
        "role": "admin, operator, technician",
        "purpose": "Menampilkan koneksi PPPoE aktif secara real-time beserta throughput per pelanggan.",
        "structure": [
            "Header dengan ringkasan: Online (badge hijau), Offline, Total Throughput Up/Down (Gbps/Mbps).",
            "Grafik area chart riwayat throughput 10 menit terakhir (data dari /api/bandwidth/traffic).",
            "Tabel sesi aktif: PPPoE Username, IP, Caller ID, Uptime, Rate Up, Rate Down, Aksi (kick).",
            "Auto-refresh setiap 30 detik.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Bandwidth Monitor                  [Refresh: 30s] |
|        +---------------------------------------------------+
|        | [Online: 131] [Offline: 17] [Tot Up: 87 Mbps]     |
|        +---------------------------------------------------+
| (sb)   | (Area Chart - Throughput 10 menit terakhir)       |
|        | ___/\___/^^\___/^\___                             |
|        +---------------------------------------------------+
|        | PPPoE       IP           Uptime   Up      Down    |
|        | budi-net    10.10.30.5   2h 15m   1.2Mbps 8.4Mbps |
|        | ani-fiber   10.10.30.7   8h 02m   0.4Mbps 12 Mbps |
|        | ...                                               |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.10 Halaman Activity Log",
        "url": "/activity-logs",
        "role": "admin, operator",
        "purpose": "Riwayat sesi login/logout PPPoE per pelanggan beserta volume traffic.",
        "structure": [
            "Filter: pelanggan, rentang tanggal, jenis aksi.",
            "Tabel: Waktu, Pelanggan, Aksi (login/logout), Durasi, Bytes In, Bytes Out.",
            "Tombol Export CSV.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Activity Log    [Plgn v] [Tgl] [Aksi v] [Export]  |
|        +---------------------------------------------------+
|        | Waktu              Plgn       Aksi    Dur. In/Out |
| (sb)   | 2026-05-08 06:01   Pak Budi   login   -    -      |
|        | 2026-05-08 14:30   Pak Budi   logout  8h   12GB.. |
|        | 2026-05-08 06:05   Bu Ani     login   -    -      |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.11 Halaman Laporan Keuangan",
        "url": "/reports",
        "role": "admin (only)",
        "purpose": "Menyajikan analitik keuangan: pendapatan, piutang, akuisisi pelanggan, churn.",
        "structure": [
            "Filter rentang waktu (bulan ini, kuartal ini, tahun ini, custom).",
            "KPI cards: Total Pendapatan, Piutang Berjalan, Pelanggan Baru, Churn Rate.",
            "Chart: pendapatan bulanan (BarChart), pertumbuhan pelanggan (LineChart), distribusi metode bayar (PieChart).",
            "Tabel ringkas: Top 10 pelanggan dengan tunggakan terlama.",
            "Tombol Cetak / Export PDF.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Laporan Keuangan   [Periode: Mei 2026 v] [Cetak]  |
|        +---------------------------------------------------+
|        | [Pendapatan] [Piutang] [Plgn Baru] [Churn]        |
|        | Rp 24.6 jt   Rp 1.8 jt   12         2.1%          |
| (sb)   |                                                   |
|        | (Bar Chart Pendapatan 12 Bulan)                   |
|        | ###  ##  ###   ##  ###  ####  ###  ##  ###  #### |
|        |                                                   |
|        | (Line Chart Pertumbuhan)  (Pie Metode Bayar)      |
|        |                                                   |
|        | -- Top 10 Pelanggan Tunggakan --                  |
|        | YBY-2412-0007 Pak Joko  3 bulan  Rp 600.000       |
+--------+---------------------------------------------------+""",
        "fields": [],
    },
    {
        "title": "3.5.12 Halaman Pengaturan dan Manajemen Pengguna",
        "url": "/settings, /settings/users",
        "role": "admin (only)",
        "purpose": "Konfigurasi koneksi MikroTik, gateway WhatsApp Fonnte, dan manajemen akun pengguna sistem.",
        "structure": [
            "Tab: Koneksi MikroTik | WhatsApp (Fonnte) | Pengguna.",
            "Tab MikroTik: Host, Port, Username, Password, tombol 'Tes Koneksi'.",
            "Tab WhatsApp: Token Fonnte, API URL, tombol 'Tes Pengiriman'.",
            "Tab Pengguna: tabel daftar staf, tombol Tambah/Edit/Hapus, form berisi nama, email, password, role.",
        ],
        "ascii": """
+--------+---------------------------------------------------+
|        | Pengaturan Sistem                                 |
|        +---------------------------------------------------+
|        | [ MikroTik ] [ WhatsApp ] [ Pengguna ]            |
|        +---------------------------------------------------+
|        | Host*       [10.10.10.1____________]              |
| (sb)   | Port*       [8728]                                |
|        | Username*   [admin___________]                    |
|        | Password    [********________]                    |
|        |                                                   |
|        | [Tes Koneksi]                  [Simpan]           |
+--------+---------------------------------------------------+""",
        "fields": [
            ["mikrotikHost", "text", "Wajib", "updateSettingsSchema"],
            ["mikrotikPort", "number", "Wajib, 1–65535, default 8728", "updateSettingsSchema"],
            ["mikrotikUser", "text", "Wajib", "updateSettingsSchema"],
            ["mikrotikPassword", "password", "Opsional", "updateSettingsSchema"],
            ["fonnteToken", "password", "Opsional", "updateSettingsSchema"],
            ["fonnteApiUrl", "text", "Opsional, default https://api.fonnte.com/send", "updateSettingsSchema"],
            ["[users.name]", "text", "Wajib, min 2 karakter", "createUserSchema"],
            ["[users.email]", "email", "Wajib, format email valid", "createUserSchema"],
            ["[users.password]", "password", "Wajib, min 8 karakter", "createUserSchema"],
            ["[users.role]", "select", "Wajib: admin/operator/technician", "createUserSchema"],
        ],
    },
]

for idx, sc in enumerate(screens):
    h(sc["title"], level=3)
    p(f"URL Rute   : {sc['url']}", indent_first=False, align="left")
    p(f"Hak Akses  : {sc['role']}", indent_first=False, align="left")
    p(f"Tujuan     : {sc['purpose']}", indent_first=False, align="left")
    p("Struktur Tata Letak:", bold=True, indent_first=False, align="left")
    for s in sc["structure"]:
        bullet(s)
    p("Wireframe:", bold=True, indent_first=False, align="left")
    code_block(sc["ascii"].strip("\n"))
    if sc["fields"]:
        p("Field & Validasi (Tabel):", bold=True, indent_first=False, align="left")
        make_table(
            headers=["Nama Field", "Tipe Input", "Aturan Validasi", "Skema Zod"],
            rows=sc["fields"],
            col_widths_cm=[3.5, 2.5, 5.5, 3],
        )

# Penutup
h("3.6 Ringkasan", level=2)
p(
    "Bab III telah memaparkan landasan teori (3.2), perancangan sistem berbasis UML mencakup use case, activity, "
    "sequence, dan class diagram (3.3), perancangan basis data PostgreSQL beserta ERD, spesifikasi tabel, dan "
    "DDL lengkap (3.4), serta perancangan layar untuk dua belas halaman utama dilengkapi struktur, wireframe, "
    "dan validasi field (3.5). Hasil perancangan ini menjadi dasar implementasi yang akan dijabarkan pada Bab IV."
)

doc.save(str(OUT))
print(f"Saved: {OUT}")
print(f"Size: {OUT.stat().st_size} bytes")
